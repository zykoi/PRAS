from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, 
                            QComboBox, QTableWidget, QTableWidgetItem, QMessageBox,
                            QFileDialog, QDialog, QTextEdit, QDateEdit, QGroupBox,
                            QCheckBox, QLineEdit, QListWidget, QListWidgetItem)
from PyQt6.QtCore import Qt, QDate
from core.database import Document, Contract, Property, Tenant, Payment
from sqlalchemy.orm import Session
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import os
import json
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication

class TemplateDialog(QDialog):
    def __init__(self, template_data=None):
        super().__init__()
        self.template_data = template_data or {}
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Редактирование шаблона")
        layout = QVBoxLayout(self)

        # Название шаблона
        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel("Название шаблона:"))
        self.name_edit = QLineEdit(self.template_data.get('name', ''))
        name_layout.addWidget(self.name_edit)
        layout.addLayout(name_layout)

        # Тип документа
        type_layout = QHBoxLayout()
        type_layout.addWidget(QLabel("Тип документа:"))
        self.type_combo = QComboBox()
        self.type_combo.addItems([
            "Договор аренды",
            "Акт приема-передачи",
            "Акт сверки",
            "Уведомление о расторжении"
        ])
        if 'type' in self.template_data:
            self.type_combo.setCurrentText(self.template_data['type'])
        type_layout.addWidget(self.type_combo)
        layout.addLayout(type_layout)

        # Содержимое шаблона
        layout.addWidget(QLabel("Содержимое шаблона:"))
        self.content_edit = QTextEdit()
        self.content_edit.setPlainText(self.template_data.get('content', ''))
        layout.addWidget(self.content_edit)

        # Кнопки
        buttons = QHBoxLayout()
        save_btn = QPushButton("Сохранить")
        save_btn.clicked.connect(self.accept)
        cancel_btn = QPushButton("Отмена")
        cancel_btn.clicked.connect(self.reject)
        buttons.addWidget(save_btn)
        buttons.addWidget(cancel_btn)
        layout.addLayout(buttons)

    def get_template_data(self):
        return {
            'name': self.name_edit.text(),
            'type': self.type_combo.currentText(),
            'content': self.content_edit.toPlainText()
        }

class EmailSettingsDialog(QDialog):
    def __init__(self, settings=None):
        super().__init__()
        self.settings = settings or {}
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Настройки email")
        layout = QVBoxLayout(self)

        # SMTP сервер
        smtp_layout = QHBoxLayout()
        smtp_layout.addWidget(QLabel("SMTP сервер:"))
        self.smtp_server = QLineEdit(self.settings.get('smtp_server', ''))
        smtp_layout.addWidget(self.smtp_server)
        layout.addLayout(smtp_layout)

        # Порт
        port_layout = QHBoxLayout()
        port_layout.addWidget(QLabel("Порт:"))
        self.port = QLineEdit(str(self.settings.get('port', '587')))
        port_layout.addWidget(self.port)
        layout.addLayout(port_layout)

        # Email
        email_layout = QHBoxLayout()
        email_layout.addWidget(QLabel("Email:"))
        self.email = QLineEdit(self.settings.get('email', ''))
        email_layout.addWidget(self.email)
        layout.addLayout(email_layout)

        # Пароль
        password_layout = QHBoxLayout()
        password_layout.addWidget(QLabel("Пароль:"))
        self.password = QLineEdit(self.settings.get('password', ''))
        self.password.setEchoMode(QLineEdit.EchoMode.Password)
        password_layout.addWidget(self.password)
        layout.addLayout(password_layout)

        # Кнопки
        buttons = QHBoxLayout()
        save_btn = QPushButton("Сохранить")
        save_btn.clicked.connect(self.accept)
        cancel_btn = QPushButton("Отмена")
        cancel_btn.clicked.connect(self.reject)
        buttons.addWidget(save_btn)
        buttons.addWidget(cancel_btn)
        layout.addLayout(buttons)

    def get_settings(self):
        return {
            'smtp_server': self.smtp_server.text(),
            'port': int(self.port.text()),
            'email': self.email.text(),
            'password': self.password.text()
        }

class BulkGenerateDialog(QDialog):
    def __init__(self, session: Session):
        super().__init__()
        self.session = session
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Массовая генерация документов")
        layout = QVBoxLayout(self)

        # Выбор типа документа
        type_layout = QHBoxLayout()
        type_layout.addWidget(QLabel("Тип документа:"))
        self.doc_type = QComboBox()
        self.doc_type.addItems([
            "Договор аренды",
            "Акт приема-передачи",
            "Акт сверки",
            "Уведомление о расторжении"
        ])
        type_layout.addWidget(self.doc_type)
        layout.addLayout(type_layout)

        # Выбор шаблона
        template_layout = QHBoxLayout()
        template_layout.addWidget(QLabel("Шаблон:"))
        self.template_combo = QComboBox()
        self.template_combo.addItem("Стандартный шаблон", None)
        template_layout.addWidget(self.template_combo)
        layout.addLayout(template_layout)

        # Список договоров
        layout.addWidget(QLabel("Выберите договоры:"))
        self.contracts_list = QListWidget()
        contracts = self.session.query(Contract).all()
        for contract in contracts:
            item = QListWidgetItem(f"Договор №{contract.id} - {contract.tenant.name if contract.tenant else '—'}")
            item.setData(Qt.ItemDataRole.UserRole, contract.id)
            item.setFlags(item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
            item.setCheckState(Qt.CheckState.Unchecked)
            self.contracts_list.addItem(item)
        layout.addWidget(self.contracts_list)

        # Опции
        options_group = QGroupBox("Опции")
        options_layout = QVBoxLayout()
        self.export_pdf = QCheckBox("Экспорт в PDF")
        self.export_pdf.setChecked(True)
        self.send_email = QCheckBox("Отправить по email")
        options_layout.addWidget(self.export_pdf)
        options_layout.addWidget(self.send_email)
        options_group.setLayout(options_layout)
        layout.addWidget(options_group)

        # Кнопки
        buttons = QHBoxLayout()
        generate_btn = QPushButton("Сформировать")
        generate_btn.clicked.connect(self.accept)
        cancel_btn = QPushButton("Отмена")
        cancel_btn.clicked.connect(self.reject)
        buttons.addWidget(generate_btn)
        buttons.addWidget(cancel_btn)
        layout.addLayout(buttons)

    def get_selected_contracts(self):
        return [
            self.contracts_list.item(i).data(Qt.ItemDataRole.UserRole)
            for i in range(self.contracts_list.count())
            if self.contracts_list.item(i).checkState() == Qt.CheckState.Checked
        ]

class DocumentsWidget(QWidget):
    def __init__(self, session: Session):
        super().__init__()
        self.session = session
        self.templates = self.load_templates()
        self.email_settings = self.load_email_settings()
        self.init_ui()

    def init_ui(self):
        # Apply stylesheet from PaymentsWidget
        self.setStyleSheet("""
            QWidget {
                background-color: #2b2b2b;
                color: #ffffff;
            }
            QPushButton {
                background-color: #0d47a1;
                color: white;
                border: none;
                padding: 8px 15px;
                border-radius: 4px;
                font-weight: bold;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #1565c0;
            }
            QPushButton:checked {
                background-color: #0a3d91;
            }
            QLabel {
                color: #ffffff;
            }
            QComboBox {
                background-color: #2b2b2b;
                color: #ffffff;
                border: 1px solid #3d3d3d;
                border-radius: 4px;
                padding: 5px;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: darkgray;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
            }
            QComboBox::down-arrow {
                image: url(icons/down_arrow.png); /* placeholder */
            }
            QGroupBox {
                color: #ffffff;
                border: 1px solid #3d3d3d;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }
            QScrollArea {
                border: none;
                background-color: transparent; /* Это важно для согласованного фона скролл-областей */
            }
            QTableWidget {
                background-color: #2b2b2b;
                color: #ffffff;
                gridline-color: #3d3d3d;
                border: none;
                border-radius: 5px;
            }
            QTableWidget::item {
                padding: 8px;
                background-color: #2b2b2b; /* Фон обычной строки */
                color: #ffffff;
            }
            QTableWidget::item:selected {
                background-color: #3d3d3d; /* Фон выбранной строки */
            }
            QHeaderView::section {
                background-color: #2b2b2b;
                color: #ffffff;
                padding: 8px;
                border: none;
                border-right: 1px solid #3d3d3d;
                border-bottom: 1px solid #3d3d3d;
                font-weight: bold;
            }
        """)

        # Apply dark theme stylesheet
        self.setStyleSheet("""
            QWidget {
                background-color: #2b2b2b;
                color: #ffffff;
            }
            /* Добавьте остальные стили по необходимости */
        """)

        layout = QVBoxLayout(self)

        # Верхняя панель с элементами управления
        controls = QHBoxLayout()
        
        # Выбор типа документа
        self.doc_type = QComboBox()
        self.doc_type.addItems([
            "Договор аренды",
            "Акт приема-передачи",
            "Акт сверки",
            "Уведомление о расторжении"
        ])
        controls.addWidget(QLabel("Тип документа:"))
        controls.addWidget(self.doc_type)

        # Выбор шаблона
        self.template_combo = QComboBox()
        self.update_templates_list()
        controls.addWidget(QLabel("Шаблон:"))
        controls.addWidget(self.template_combo)

        # Выбор договора
        self.contract_combo = QComboBox()
        contracts = self.session.query(Contract).all()
        for contract in contracts:
            self.contract_combo.addItem(
                f"Договор №{contract.id} - {contract.tenant.name if contract.tenant else '—'}",
                contract.id
            )
        controls.addWidget(QLabel("Договор:"))
        controls.addWidget(self.contract_combo)

        layout.addLayout(controls)

        # Новый горизонтальный макет для кнопок
        button_layout = QHBoxLayout()

        # Кнопки
        generate_btn = QPushButton("Сформировать документ")
        generate_btn.setMinimumWidth(180)
        generate_btn.clicked.connect(self.generate_document)
        button_layout.addWidget(generate_btn)

        bulk_generate_btn = QPushButton("Массовая генерация")
        bulk_generate_btn.setMinimumWidth(180)
        bulk_generate_btn.clicked.connect(self.show_bulk_generate_dialog)
        button_layout.addWidget(bulk_generate_btn)

        manage_templates_btn = QPushButton("Управление шаблонами")
        manage_templates_btn.setMinimumWidth(180)
        manage_templates_btn.clicked.connect(self.manage_templates)
        button_layout.addWidget(manage_templates_btn)

        email_settings_btn = QPushButton("Настройки email")
        email_settings_btn.setMinimumWidth(180)
        email_settings_btn.clicked.connect(self.show_email_settings)
        button_layout.addWidget(email_settings_btn)

        # Добавляем макет с кнопками в основной макет
        layout.addLayout(button_layout)

        # Опции экспорта
        export_group = QGroupBox("Опции экспорта")
        export_layout = QHBoxLayout()
        self.export_pdf = QCheckBox("Экспорт в PDF")
        self.export_pdf.setChecked(True)
        self.send_email = QCheckBox("Отправить по email")
        export_layout.addWidget(self.export_pdf)
        export_layout.addWidget(self.send_email)
        export_group.setLayout(export_layout)
        layout.addWidget(export_group)

        # Таблица с историей документов
        self.table = QTableWidget()
        self.table.setColumnCount(7)
        self.table.setHorizontalHeaderLabels([
            "Дата", "Тип документа", "Договор", "Арендатор", "Файл", "PDF", "Email"
        ])
        layout.addWidget(self.table)

        # Загружаем историю документов
        self.load_documents_history()

    def load_templates(self):
        try:
            with open('templates.json', 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            return []

    def save_templates(self):
        with open('templates.json', 'w', encoding='utf-8') as f:
            json.dump(self.templates, f, ensure_ascii=False, indent=2)

    def update_templates_list(self):
        self.template_combo.clear()
        self.template_combo.addItem("Стандартный шаблон", None)
        for template in self.templates:
            if template['type'] == self.doc_type.currentText():
                self.template_combo.addItem(template['name'], template)

    def manage_templates(self):
        dialog = TemplateDialog()
        if dialog.exec():
            template_data = dialog.get_template_data()
            self.templates.append(template_data)
            self.save_templates()
            self.update_templates_list()

    def generate_document(self):
        doc_type = self.doc_type.currentText()
        contract_id = self.contract_combo.currentData()
        contract = self.session.query(Contract).get(contract_id)

        if not contract:
            QMessageBox.warning(self, "Ошибка", "Договор не найден")
            return

        # Запрашиваем место сохранения
        file_name, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить документ",
            f"{doc_type}_{contract.id}_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )

        if file_name:
            # Получаем выбранный шаблон
            template = self.template_combo.currentData()
            
            if template:
                # Используем пользовательский шаблон
                doc = self.generate_from_template(template, contract)
            else:
                # Используем стандартный шаблон
                if doc_type == "Договор аренды":
                    doc = self.generate_contract(contract)
                elif doc_type == "Акт приема-передачи":
                    doc = self.generate_handover_act(contract)
                elif doc_type == "Акт сверки":
                    doc = self.generate_reconciliation_act(contract)
                elif doc_type == "Уведомление о расторжении":
                    doc = self.generate_termination_notice(contract)

            # Сохраняем документ
            doc.save(file_name)

            # Экспортируем в PDF если выбрано
            if self.export_pdf.isChecked():
                pdf_file = os.path.splitext(file_name)[0] + '.pdf'
                convert(file_name, pdf_file)
                QMessageBox.information(self, "Успех", 
                    f"Документ успешно сформирован\nWord: {file_name}\nPDF: {pdf_file}")
            else:
                QMessageBox.information(self, "Успех", "Документ успешно сформирован")

    def generate_from_template(self, template, contract):
        doc = Document()
        content = template['content']
        
        # Заменяем плейсхолдеры на реальные данные
        content = content.replace('{contract_id}', str(contract.id))
        content = content.replace('{start_date}', contract.start_date.strftime("%d.%m.%Y"))
        content = content.replace('{end_date}', contract.end_date.strftime("%d.%m.%Y"))
        content = content.replace('{tenant_name}', contract.tenant.name if contract.tenant else '—')
        content = content.replace('{tenant_inn}', contract.tenant.legal_info)
        content = content.replace('{tenant_address}', contract.tenant.contact_info)
        content = content.replace('{property_address}', contract.property.address)
        content = content.replace('{area}', str(contract.area))
        content = content.replace('{monthly_rent}', f"{contract.monthly_rent:.2f}")

        # Добавляем содержимое в документ
        for paragraph in content.split('\n'):
            doc.add_paragraph(paragraph)

        return doc

    def generate_contract(self, contract):
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('ДОГОВОР АРЕНДЫ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Номер и дата
        doc.add_paragraph(f'№ {contract.id} от {contract.start_date.strftime("%d.%m.%Y")}')
        
        # Стороны договора
        doc.add_paragraph('\nАРЕНДОДАТЕЛЬ:')
        doc.add_paragraph('ООО "РентКом"')
        doc.add_paragraph('ИНН: 1234567890')
        doc.add_paragraph('Адрес: г. Москва, ул. Примерная, д. 1')
        
        doc.add_paragraph('\nАРЕНДАТОР:')
        doc.add_paragraph(contract.tenant.name if contract.tenant else '—')
        doc.add_paragraph(f'ИНН: {contract.tenant.legal_info}')
        doc.add_paragraph(f'Адрес: {contract.tenant.contact_info}')

        # Предмет договора
        doc.add_paragraph('\n1. ПРЕДМЕТ ДОГОВОРА')
        doc.add_paragraph(f'1.1. Арендодатель передает, а Арендатор принимает в аренду помещение:')
        doc.add_paragraph(f'Адрес: {contract.property.address}')
        doc.add_paragraph(f'Площадь: {contract.area} кв.м')
        
        # Срок аренды
        doc.add_paragraph('\n2. СРОК АРЕНДЫ')
        doc.add_paragraph(f'2.1. Срок аренды: с {contract.start_date.strftime("%d.%m.%Y")} по {contract.end_date.strftime("%d.%m.%Y")}')

        # Арендная плата
        doc.add_paragraph('\n3. АРЕНДНАЯ ПЛАТА')
        doc.add_paragraph(f'3.1. Размер арендной платы: {contract.monthly_rent:.2f} рублей в месяц')
        doc.add_paragraph('3.2. Арендная плата вносится ежемесячно до 5 числа текущего месяца')

        return doc

    def generate_handover_act(self, contract):
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('АКТ ПРИЕМА-ПЕРЕДАЧИ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Дата
        doc.add_paragraph(f'от {datetime.now().strftime("%d.%m.%Y")}')
        
        # Стороны
        doc.add_paragraph('\nАРЕНДОДАТЕЛЬ:')
        doc.add_paragraph('ООО "РентКом"')
        
        doc.add_paragraph('\nАРЕНДАТОР:')
        doc.add_paragraph(contract.tenant.name if contract.tenant else '—')

        # Описание помещения
        doc.add_paragraph('\nПомещение передано в аренду:')
        doc.add_paragraph(f'Адрес: {contract.property.address}')
        doc.add_paragraph(f'Площадь: {contract.area} кв.м')
        
        # Состояние помещения
        doc.add_paragraph('\nСостояние помещения:')
        doc.add_paragraph('Помещение передано в исправном состоянии')
        
        # Подписи
        doc.add_paragraph('\nАрендодатель: _________________')
        doc.add_paragraph('\nАрендатор: _________________')

        return doc

    def generate_reconciliation_act(self, contract):
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('АКТ СВЕРКИ РАСЧЕТОВ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Дата
        doc.add_paragraph(f'от {datetime.now().strftime("%d.%m.%Y")}')
        
        # Стороны
        doc.add_paragraph('\nАРЕНДОДАТЕЛЬ:')
        doc.add_paragraph('ООО "РентКом"')
        
        doc.add_paragraph('\nАРЕНДАТОР:')
        doc.add_paragraph(contract.tenant.name if contract.tenant else '—')

        # Расчеты
        doc.add_paragraph('\nРасчеты по договору аренды:')
        payments = self.session.query(Payment).filter(
            Payment.contract_id == contract.id
        ).all()
        
        total_debt = 0
        for payment in payments:
            if payment.status == 'pending' or payment.status == 'overdue':
                total_debt += payment.amount

        doc.add_paragraph(f'Сумма задолженности: {total_debt:.2f} рублей')
        
        # Подписи
        doc.add_paragraph('\nАрендодатель: _________________')
        doc.add_paragraph('\nАрендатор: _________________')

        return doc

    def generate_termination_notice(self, contract):
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('УВЕДОМЛЕНИЕ О РАСТОРЖЕНИИ ДОГОВОРА', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Дата
        doc.add_paragraph(f'от {datetime.now().strftime("%d.%m.%Y")}')
        
        # Адресат
        doc.add_paragraph('\nАРЕНДАТОРУ:')
        doc.add_paragraph(contract.tenant.name if contract.tenant else '—')
        doc.add_paragraph(contract.tenant.contact_info)

        # Текст уведомления
        doc.add_paragraph('\nНастоящим уведомляем Вас о расторжении договора аренды №{contract.id} от {contract.start_date.strftime("%d.%m.%Y")}.')
        doc.add_paragraph('Договор считается расторгнутым с момента получения настоящего уведомления.')
        doc.add_paragraph('Просим Вас освободить помещение и передать его по акту приема-передачи в течение 30 дней с момента получения настоящего уведомления.')

        # Подписи
        doc.add_paragraph('\nАрендодатель: _________________')

        return doc

    def load_documents_history(self):
        # В реальном приложении здесь будет загрузка из базы данных
        # Сейчас просто заглушка
        self.table.setRowCount(0)

    def load_email_settings(self):
        try:
            with open('email_settings.json', 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            return {}

    def save_email_settings(self):
        with open('email_settings.json', 'w', encoding='utf-8') as f:
            json.dump(self.email_settings, f, ensure_ascii=False, indent=2)

    def show_email_settings(self):
        dialog = EmailSettingsDialog(self.email_settings)
        if dialog.exec():
            self.email_settings = dialog.get_settings()
            self.save_email_settings()

    def show_bulk_generate_dialog(self):
        dialog = BulkGenerateDialog(self.session)
        if dialog.exec():
            doc_type = dialog.doc_type.currentText()
            template = dialog.template_combo.currentData()
            contract_ids = dialog.get_selected_contracts()
            export_pdf = dialog.export_pdf.isChecked()
            send_email = dialog.send_email.isChecked()

            if not contract_ids:
                QMessageBox.warning(self, "Ошибка", "Выберите хотя бы один договор")
                return

            # Создаем папку для документов
            folder_name = f"documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            os.makedirs(folder_name, exist_ok=True)

            # Генерируем документы
            for contract_id in contract_ids:
                contract = self.session.query(Contract).get(contract_id)
                if not contract:
                    continue

                # Генерируем документ
                file_name = os.path.join(folder_name, f"{doc_type}_{contract.id}_{datetime.now().strftime('%Y%m%d')}.docx")
                
                if template:
                    doc = self.generate_from_template(template, contract)
                else:
                    if doc_type == "Договор аренды":
                        doc = self.generate_contract(contract)
                    elif doc_type == "Акт приема-передачи":
                        doc = self.generate_handover_act(contract)
                    elif doc_type == "Акт сверки":
                        doc = self.generate_reconciliation_act(contract)
                    elif doc_type == "Уведомление о расторжении":
                        doc = self.generate_termination_notice(contract)

                doc.save(file_name)

                # Экспортируем в PDF если выбрано
                if export_pdf:
                    pdf_file = os.path.splitext(file_name)[0] + '.pdf'
                    convert(file_name, pdf_file)

                # Отправляем по email если выбрано
                if send_email and self.email_settings:
                    self.send_document_by_email(contract, file_name, doc_type)

            QMessageBox.information(self, "Успех", f"Документы сформированы в папке {folder_name}")

    def send_document_by_email(self, contract, file_name, doc_type):
        if not self.email_settings:
            return

        try:
            # Создаем сообщение
            msg = MIMEMultipart()
            msg['From'] = self.email_settings['email']
            msg['To'] = contract.tenant.contact_info
            msg['Subject'] = f"{doc_type} - Договор №{contract.id}"

            # Добавляем текст письма
            body = f"""
            Здравствуйте!

            В приложении находится {doc_type} по договору №{contract.id}.

            С уважением,
            ООО "РентКом"
            """
            msg.attach(MIMEText(body, 'plain'))

            # Добавляем вложения
            with open(file_name, 'rb') as f:
                part = MIMEApplication(f.read(), Name=os.path.basename(file_name))
                part['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_name)}"'
                msg.attach(part)

            # Добавляем PDF если есть
            pdf_file = os.path.splitext(file_name)[0] + '.pdf'
            if os.path.exists(pdf_file):
                with open(pdf_file, 'rb') as f:
                    part = MIMEApplication(f.read(), Name=os.path.basename(pdf_file))
                    part['Content-Disposition'] = f'attachment; filename="{os.path.basename(pdf_file)}"'
                    msg.attach(part)

            # Отправляем письмо
            with smtplib.SMTP(self.email_settings['smtp_server'], self.email_settings['port']) as server:
                server.starttls()
                server.login(self.email_settings['email'], self.email_settings['password'])
                server.send_message(msg)

            return True
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось отправить email: {str(e)}")
            return False 